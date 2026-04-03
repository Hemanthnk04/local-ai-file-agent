"""
document_tools.py — read/write Word, PDF, and Excel files.

All heavy libraries (docx, pypdf, reportlab, openpyxl) are imported
LAZILY inside each function so they only load when actually needed.
This keeps startup fast and RAM low on small machines.
"""


# ── WORD ──────────────────────────────────────────────────────────────────────

def create_word(path, text):
    from docx import Document
    doc = Document()
    for line in str(text).split("\n"):
        doc.add_paragraph(line)
    doc.save(path)


def read_word(path):
    from docx import Document
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


# ── PDF ───────────────────────────────────────────────────────────────────────

def create_pdf(path, text):
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas
    c = rl_canvas.Canvas(path, pagesize=A4)
    y = 820
    for line in str(text).split("\n"):
        c.drawString(50, y, line[:120])
        y -= 16
        if y < 50:
            c.showPage()
            y = 820
    c.save()


def read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts  = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    return "\n".join(parts)


# ── EXCEL single sheet ────────────────────────────────────────────────────────

def create_excel(path, rows):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    for r in rows:
        ws.append(r)
    wb.save(path)


def read_excel(path):
    from openpyxl import load_workbook
    wb   = load_workbook(path, read_only=True, data_only=True)
    ws   = wb.active
    data = [list(row) for row in ws.iter_rows(values_only=True)]
    wb.close()
    return data


# ── EXCEL multiple sheets ─────────────────────────────────────────────────────

def create_excel_multi(path, sheets):
    """
    sheets: {sheet_name: [[row], [row], ...]}
    Delegates to the robust build_excel_multisheet in content_validator.
    """
    from .content_validator import build_excel_multisheet
    ok, reason = build_excel_multisheet(sheets, path)
    if not ok:
        raise RuntimeError(f"Excel multi-sheet write failed: {reason}")
