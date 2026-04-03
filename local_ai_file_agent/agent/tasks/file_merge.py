"""
FILE_MERGE — merge any number of files of any supported type.

Strategy by category:
  Data (csv/tsv/xlsx/json/xml/yaml) → stack rows into one file, OR
                                       multi-sheet Excel if user asks
  Code / text / config / html / md  → LLM-assisted merge preserving all logic
  PDF / DOCX                         → concatenate text content
  Mixed types                        → ask user for target format, then convert+merge
"""

import os
import re
import json
import glob
from ..resolve import resolve_multi
from ..guardrails import run_write_guardrails
from ..utils import confirm
from ..config import BASE_DIR
from ..llm import call_llm
from ..prompts import CODE_LANG


# ── Extension categories ──────────────────────────────────────────────────────
DATA_EXTS   = {".csv", ".tsv", ".xlsx", ".json", ".xml", ".yaml", ".yml"}
CODE_EXTS   = {".py", ".pyw", ".js", ".ts", ".jsx", ".tsx", ".java", ".sql",
               ".css", ".cpp", ".c", ".cs", ".go", ".rb", ".php"}
TEXT_EXTS   = {".txt", ".md", ".markdown", ".html", ".htm",
               ".toml", ".ini", ".cfg", ".env"}
DOC_EXTS    = {".pdf", ".docx"}

ALL_MERGE_EXTS = DATA_EXTS | CODE_EXTS | TEXT_EXTS | DOC_EXTS


# ── File collection ───────────────────────────────────────────────────────────

def _collect_files(filename, instructions):
    """Collect source files from classifier output, glob, or interactive picker."""
    sources = []

    # 1. Glob pattern in instructions
    if instructions:
        m = re.search(r"[\w./\\*]+\*[\w./\\*]*", instructions)
        if m:
            hits = glob.glob(m.group(0).strip(), recursive=True)
            if hits:
                return sorted(hits)

    # 2. Comma-separated filenames from classifier
    if filename:
        sources = resolve_multi(filename, instructions)
        if len(sources) >= 2:
            return sources

    # 3. Interactive — ask for folder and extension
    print("\n  No source files identified from your prompt.")
    folder = input("  Folder containing files to merge (Enter = current dir): ").strip() or BASE_DIR

    print("  Supported: csv, tsv, xlsx, json, xml, yaml, py, js, ts, java, sql,")
    print("             cpp, c, cs, go, rb, php, css, html, md, txt, toml, ini, pdf, docx")
    ext_raw = input("  File extension to merge (e.g. csv, py, js): ").strip().lower().lstrip(".")
    if not ext_raw:
        return []

    ext  = "." + ext_raw
    hits = sorted(glob.glob(os.path.join(folder, f"*{ext}")))

    if not hits:
        print(f"  No {ext} files found in '{folder}'.")
        return []

    print(f"\n  Found {len(hits)} {ext} file(s):")
    for i, h in enumerate(hits):
        print(f"    [{i}] {os.path.basename(h)}")

    raw = input("  Select numbers (comma-separated, Enter = all): ").strip()
    if not raw:
        return hits
    try:
        indices = [int(x.strip()) for x in raw.split(",")]
        return [hits[i] for i in indices if 0 <= i < len(hits)]
    except ValueError:
        return hits


# ── Readers ───────────────────────────────────────────────────────────────────

def _read_as_text(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def _read_docx_text(path):
    from docx import Document
    return "\n".join(p.text for p in Document(path).paragraphs)

def _read_pdf_text(path):
    from pypdf import PdfReader
    return "\n\n".join(
        p.extract_text() for p in PdfReader(path).pages if p.extract_text()
    )

def _read_to_df(path, ext):
    import pandas as pd
    if ext == ".csv":  return pd.read_csv(path)
    if ext == ".tsv":  return pd.read_csv(path, sep="\t")
    if ext == ".xlsx": return pd.read_excel(path)
    if ext == ".json":
        with open(path) as f:
            data = json.load(f)
        if isinstance(data, list): return pd.DataFrame(data)
        if isinstance(data, dict): return pd.DataFrame([data])
    if ext == ".xml":
        import xml.etree.ElementTree as ET
        root = ET.parse(path).getroot()
        rows = [{c.tag: c.text for c in child} for child in root]
        return pd.DataFrame(rows)
    if ext in (".yaml", ".yml"):
        import yaml
        with open(path) as f:
            data = yaml.safe_load(f)
        if isinstance(data, list): return pd.DataFrame(data)
        if isinstance(data, dict): return pd.DataFrame([data])
    return None


# ── Writers ───────────────────────────────────────────────────────────────────

def _write_merged_data(dfs, names, out_path, ext, multi_sheet, join_how):
    import pandas as pd
    """Write merged data files. Multi-sheet → always .xlsx."""
    if multi_sheet:
        from ..content_validator import build_excel_multisheet
        actual_out   = os.path.splitext(out_path)[0] + ".xlsx"
        sheet_data   = {}
        for name, df in zip(names, dfs):
            sheet_name = os.path.splitext(os.path.basename(name))[0]
            sheet_data[sheet_name] = df
        ok, reason = build_excel_multisheet(sheet_data, actual_out)
        if not ok:
            return actual_out, f"Multi-sheet error: {reason}"
        return actual_out, f"{len(dfs)} sheets written ({reason})"

    # Single merged file
    merged = pd.concat(dfs, join=join_how, ignore_index=True)

    if ext in (".csv", ".tsv"):
        sep = "\t" if ext == ".tsv" else ","
        merged.to_csv(out_path, index=False, sep=sep)
    elif ext == ".xlsx":
        merged.to_excel(out_path, index=False)
    elif ext == ".json":
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(merged.to_dict(orient="records"), f, indent=2, ensure_ascii=False)
    elif ext == ".xml":
        import xml.etree.ElementTree as ET
        root = ET.Element("data")
        for _, row in merged.iterrows():
            item = ET.SubElement(root, "row")
            for col, val in row.items():
                c = ET.SubElement(item, str(col).replace(" ", "_"))
                c.text = str(val)
        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ")
        tree.write(out_path, encoding="unicode", xml_declaration=True)
    elif ext in (".yaml", ".yml"):
        import yaml
        with open(out_path, "w", encoding="utf-8") as f:
            yaml.dump(merged.to_dict(orient="records"), f,
                      allow_unicode=True, default_flow_style=False)
    else:
        merged.to_csv(out_path, index=False)

    return out_path, f"{len(merged):,} rows, {len(merged.columns)} columns"


# Comment styles for file separators by extension
_COMMENT_STYLE = {
    ".py": "# ", ".pyw": "# ", ".rb": "# ", ".sh": "# ",
    ".bash": "# ", ".yaml": "# ", ".yml": "# ", ".toml": "# ",
    ".sql": "-- ", ".md": "<!-- ", ".html": "<!-- ", ".htm": "<!-- ",
    ".css": "/* ", ".c": "// ", ".cpp": "// ", ".cs": "// ",
    ".go": "// ", ".java": "// ", ".js": "// ", ".ts": "// ",
    ".jsx": "// ", ".tsx": "// ", ".php": "// ",
}
_COMMENT_CLOSE = {".css": " */", ".md": " -->", ".html": " -->", ".htm": " -->"}


def _write_merged_text(contents, names, out_path, ext, instructions):
    """Write merged text/code files — concat with language-appropriate separators."""
    prefix = _COMMENT_STYLE.get(ext, "")
    suffix = _COMMENT_CLOSE.get(ext, "")

    parts = []
    for name, content in zip(names, contents):
        basename = os.path.basename(name)
        _dash = "─" * max(0, 48 - len(basename))
        sep_line = f"{prefix}── {basename} {_dash}{suffix}"
        parts.append(f"{sep_line}\n{content}")

    combined = "\n\n".join(parts)
    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(combined)
    return out_path, f"{len(contents)} files concatenated"


def _write_merged_code_via_llm(contents, names, out_path, ext, instructions):
    """
    Use the LLM to merge code files intelligently:
    send all files with clear FILE: markers and ask it to merge without losing logic.
    """
    lang = CODE_LANG.get(ext, ext.lstrip("."))

    file_blocks = "\n\n".join(
        f"FILE: {os.path.basename(n)}\n{c}"
        for n, c in zip(names, contents)
    )

    prompt = f"""You are merging multiple {lang} source files into one.

RULES:
1. Preserve ALL functions, classes, imports, and logic from EVERY file.
2. Resolve naming conflicts by keeping both with a comment noting the origin file.
3. Remove duplicate imports — keep only one copy.
4. Do NOT summarise, truncate, or omit any code.
5. Output ONLY the merged {lang} code — no explanations, no markdown fences.
6. Add a short comment at the top listing the source files merged.
7. Maintain proper {lang} syntax throughout.

USER INSTRUCTION: {instructions or f'Merge all {lang} files into one cohesive file.'}

SOURCE FILES:
{file_blocks}

OUTPUT THE COMPLETE MERGED {lang.upper()} CODE ONLY.
"""
    result = call_llm(prompt)
    # Strip any accidental markdown fences
    result = re.sub(r"^```[\w]*\n?", "", result, flags=re.M).strip()
    result = re.sub(r"\n?```$", "", result, flags=re.M).strip()

    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(result)
    return out_path, f"{len(names)} files merged via LLM"


def _write_merged_docs(contents, names, out_path, ext):
    import pandas as pd
    """Concatenate text documents with clear section headers."""
    parts = []
    for name, content in zip(names, contents):
        parts.append(f"{'═'*60}\nSOURCE: {os.path.basename(name)}\n{'═'*60}\n{content}")
    combined = "\n\n".join(parts)

    if ext == ".txt" or ext == ".md":
        with open(out_path, "w", encoding="utf-8", newline="\n") as f:
            f.write(combined)
    elif ext == ".docx":
        from docx import Document
        doc = Document()
        for name, content in zip(names, contents):
            doc.add_heading(os.path.basename(name), level=1)
            for line in content.splitlines():
                doc.add_paragraph(line)
        doc.save(out_path)
    elif ext == ".pdf":
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch
        styles = getSampleStyleSheet()
        story  = []
        for name, content in zip(names, contents):
            story.append(Paragraph(f"<b>{os.path.basename(name)}</b>", styles["Heading1"]))
            for line in content.splitlines():
                if line.strip():
                    story.append(Paragraph(
                        line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"),
                        styles["Normal"]
                    ))
                else:
                    story.append(Spacer(1, 0.1*inch))
        SimpleDocTemplate(out_path, pagesize=A4).build(story)
    else:
        with open(out_path, "w", encoding="utf-8", newline="\n") as f:
            f.write(combined)

    return out_path, f"{len(contents)} documents merged"


# ── Entry point ───────────────────────────────────────────────────────────────

def run(filename=None, instructions=None, **kwargs):

    sources = _collect_files(filename, instructions)

    if len(sources) < 2:
        print("  Need at least 2 files to merge.")
        return

    # Detect extensions
    exts      = [os.path.splitext(p)[1].lower() for p in sources]
    first_ext = exts[0]
    mixed     = len(set(exts)) > 1

    # ── Show plan ─────────────────────────────────────────────────────────
    print(f"\n  {'─'*52}")
    print(f"  MERGE PLAN")
    print(f"  {'─'*52}")
    for s in sources:
        print(f"    • {s}")

    if mixed:
        print(f"\n  ⚠  Mixed types: {sorted(set(exts))}")
        print("     All files will be read as text and concatenated.")

    # ── Multi-sheet Excel option (only for data files) ────────────────────
    multi_sheet = False
    if not mixed and first_ext in DATA_EXTS:
        low = (instructions or "").lower()
        if any(k in low for k in ("multi sheet", "multiple sheet", "each sheet",
                                   "separate sheet", "multisheet", "sheets")):
            multi_sheet = True
            print("\n  📊 Multi-sheet Excel mode detected.")
        elif first_ext not in (".xlsx",):
            choice = input("\n  Merge into: 1 = single file  2 = multi-sheet Excel: ").strip()
            if choice == "2":
                multi_sheet = True

    # ── Output filename ───────────────────────────────────────────────────
    out_ext     = ".xlsx" if multi_sheet else first_ext
    folder      = os.path.dirname(sources[0]) or BASE_DIR
    out_default = os.path.join(folder, "merged" + out_ext)
    out_path    = input(f"\n  Output filename [{out_default}]: ").strip() or out_default

    # Ensure correct extension
    if not os.path.splitext(out_path)[1]:
        out_path += out_ext

    ok, reason = run_write_guardrails(out_path)
    if not ok:
        print(f"  Blocked: {reason}")
        return

    if os.path.exists(out_path):
        if not confirm(f"  '{os.path.basename(out_path)}' already exists. Overwrite? (yes/no): "):
            return

    if not confirm(f"\n  Merge {len(sources)} files → '{os.path.basename(out_path)}'? (yes/no): "):
        print("  Cancelled.")
        return

    # ── Execute merge ─────────────────────────────────────────────────────
    try:
        if not mixed and first_ext in DATA_EXTS:
            # Data merge
            join_how = "inner" if "matching" in (instructions or "").lower() else "outer"
            dfs = []
            for p in sources:
                try:
                    df = _read_to_df(p, os.path.splitext(p)[1].lower())
                    if df is not None:
                        dfs.append(df)
                    else:
                        print(f"  ⚠  Could not read as table: {p}")
                except Exception as e:
                    print(f"  ⚠  Skipping '{p}': {e}")

            if not dfs:
                print("  ❌ No files could be read as data.")
                return

            final, summary = _write_merged_data(
                dfs, sources, out_path, first_ext, multi_sheet, join_how
            )

        elif not mixed and first_ext in CODE_EXTS:
            # Code merge via LLM
            print(f"  🤖 Using LLM to merge {CODE_LANG.get(first_ext, first_ext)} code...")
            contents = [_read_as_text(p) for p in sources]
            final, summary = _write_merged_code_via_llm(
                contents, sources, out_path, first_ext, instructions
            )

        elif not mixed and first_ext in DOC_EXTS:
            # Document merge
            contents = []
            for p in sources:
                ext = os.path.splitext(p)[1].lower()
                if ext == ".docx":   contents.append(_read_docx_text(p))
                elif ext == ".pdf":  contents.append(_read_pdf_text(p))
                else:                contents.append(_read_as_text(p))
            final, summary = _write_merged_docs(contents, sources, out_path, first_ext)

        else:
            # Text / config / mixed — concatenate with separators
            contents = []
            for p in sources:
                ext = os.path.splitext(p)[1].lower()
                if ext == ".docx":   contents.append(_read_docx_text(p))
                elif ext == ".pdf":  contents.append(_read_pdf_text(p))
                else:                contents.append(_read_as_text(p))
            final, summary = _write_merged_text(contents, sources, out_path, first_ext, instructions)

    except Exception as e:
        print(f"  ❌ Merge error: {e}")
        import traceback; traceback.print_exc()
        return

    size     = os.path.getsize(final)
    size_str = f"{size:,} B" if size < 1024 else f"{size//1024:,} KB"
    print(f"\n  ✅ Merged {len(sources)} files → {final}")
    print(f"     {summary}  ({size_str})")
