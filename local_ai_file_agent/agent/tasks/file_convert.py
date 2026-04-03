"""
FILE_CONVERT — convert between all supported file types.

Universal targets (any file → these):
    txt, md, pdf, docx, html

Data targets (data files ↔):
    csv, tsv, json, xlsx, xml, yaml

Code/config are text-family — they convert to/from txt, md, pdf, docx, html

Full supported source extensions:
    Data    : .csv .tsv .json .xlsx .xml .yaml .yml
    Document: .txt .md .pdf .docx .html .htm
    Code    : .py .js .ts .jsx .tsx .java .sql .css .cpp .c .cs .go .rb .php
    Config  : .yaml .yml .toml .ini .cfg .env
"""

import os
import json
import re
from io import StringIO
from ..resolve import resolve_single
from ..guardrails import run_write_guardrails
from ..utils import confirm


# ══════════════════════════════════════════════════════════════════════════════
# CATEGORIES
# ══════════════════════════════════════════════════════════════════════════════

DATA_EXTS   = {".csv", ".tsv", ".json", ".xlsx", ".xml", ".yaml", ".yml"}
DOC_EXTS    = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}
CODE_EXTS   = {".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".sql",
               ".css", ".cpp", ".c", ".cs", ".go", ".rb", ".php"}
CONFIG_EXTS = {".yaml", ".yml", ".toml", ".ini", ".cfg", ".env"}

ALL_READABLE = DATA_EXTS | DOC_EXTS | CODE_EXTS | CONFIG_EXTS

# Universal targets — any source can go to these
UNIVERSAL_TARGETS = {".txt", ".md", ".pdf", ".docx", ".html"}

# Data ↔ data targets
DATA_TARGETS = {".csv", ".tsv", ".json", ".xlsx", ".xml", ".yaml"}

def _allowed_targets(src_ext):
    targets = set(UNIVERSAL_TARGETS)
    if src_ext in DATA_EXTS:
        targets |= DATA_TARGETS
    targets.discard(src_ext)  # can't convert to same type
    # .htm is an alias for .html
    if src_ext == ".htm":
        targets.discard(".html")
    if src_ext == ".yml":
        targets.discard(".yaml")
    return targets


# ══════════════════════════════════════════════════════════════════════════════
# READERS  →  returns (data, kind)
# kind: "dataframe" | "text" | "json_obj" | "xml_tree" | "yaml_obj"
# ══════════════════════════════════════════════════════════════════════════════

def _read_csv(path):
    import pandas as pd
    return pd.read_csv(path), "dataframe"

def _read_tsv(path):
    import pandas as pd
    return pd.read_csv(path, sep="\t"), "dataframe"

def _read_json(path):
    import pandas as pd
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list):
        try:
            return pd.DataFrame(data), "dataframe"
        except Exception:
            pass
    return data, "json_obj"

def _read_xlsx(path):
    import pandas as pd
    return pd.read_excel(path), "dataframe"

def _read_xml(path):
    import xml.etree.ElementTree as ET
    tree = ET.parse(path)
    root = tree.getroot()
    # Try to convert to dataframe (list of same-tag children)
    children = list(root)
    if children and len(set(c.tag for c in children)) == 1:
        rows = [{**c.attrib, **{sub.tag: sub.text for sub in c}} for c in children]
        try:
            return pd.DataFrame(rows), "dataframe"
        except Exception:
            pass
    return (root, tree), "xml_tree"

def _read_yaml(path):
    import pandas as pd
    import yaml
    with open(path, "r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    if isinstance(data, list):
        try:
            return pd.DataFrame(data), "dataframe"
        except Exception:
            pass
    return data, "yaml_obj"

def _read_text(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read(), "text"

def _read_docx(path):
    from docx import Document
    doc  = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    return text, "text"

def _read_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages  = [p.extract_text() for p in reader.pages if p.extract_text()]
    return "\n\n".join(pages), "text"

def _read_html(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()
    # Strip tags for plain text extraction
    text = re.sub(r"<[^>]+>", " ", raw)
    text = re.sub(r"\s{2,}", "\n", text).strip()
    return (raw, text), "html_pair"   # (raw_html, plain_text)


READERS = {
    ".csv":  _read_csv,
    ".tsv":  _read_tsv,
    ".json": _read_json,
    ".xlsx": _read_xlsx,
    ".xml":  _read_xml,
    ".yaml": _read_yaml,
    ".yml":  _read_yaml,
    ".txt":  _read_text,
    ".md":   _read_text,
    ".pdf":  _read_pdf,
    ".docx": _read_docx,
    ".html": _read_html,
    ".htm":  _read_html,
}
# All code/config/shell extensions use _read_text
for _ext in CODE_EXTS | CONFIG_EXTS | {".bash", ".zsh", ".pyw", ".sh", ".markdown"}:
    if _ext not in READERS:
        READERS[_ext] = _read_text


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS — data → text representations
# ══════════════════════════════════════════════════════════════════════════════

def _to_plain_text(data, kind):
    if kind == "dataframe":
        return data.to_string(index=False)
    if kind in ("json_obj", "yaml_obj"):
        return json.dumps(data, indent=2, ensure_ascii=False)
    if kind == "xml_tree":
        import xml.etree.ElementTree as ET
        root, _ = data
        return ET.tostring(root, encoding="unicode")
    if kind == "html_pair":
        return data[1]   # plain text version
    return str(data)

def _to_dataframe(data, kind, src_ext):
    import pandas as pd
    """Best-effort conversion to DataFrame."""
    if kind == "dataframe":
        return data
    if kind in ("json_obj", "yaml_obj"):
        if isinstance(data, list):
            return pd.DataFrame(data)
        if isinstance(data, dict):
            return pd.DataFrame([data])
    if kind == "xml_tree":
        root, _ = data
        rows = []
        for child in root:
            row = {**child.attrib}
            for sub in child:
                row[sub.tag] = sub.text
            rows.append(row)
        return pd.DataFrame(rows)
    if kind == "text":
        try:
            sep = "\t" if src_ext == ".tsv" else ","
            return pd.read_csv(StringIO(str(data)), sep=sep)
        except Exception:
            pass
    if kind == "html_pair":
        try:
            dfs = pd.read_html(data[0])
            return dfs[0] if dfs else None
        except Exception:
            pass
    return None


# ══════════════════════════════════════════════════════════════════════════════
# WRITERS
# ══════════════════════════════════════════════════════════════════════════════

def _write_txt(data, kind, src_ext, out_path):
    text = _to_plain_text(data, kind)
    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(text)

def _write_md(data, kind, src_ext, out_path):
    import pandas as pd
    """Convert to Markdown. DataFrames → GFM table. Code → fenced block."""
    if kind == "dataframe":
        md = data.to_markdown(index=False) if hasattr(data, "to_markdown") else data.to_string(index=False)
    elif src_ext in CODE_EXTS:
        lang = src_ext.lstrip(".")
        md   = f"```{lang}\n{_to_plain_text(data, kind)}\n```"
    else:
        md = _to_plain_text(data, kind)
    with open(out_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(md)

def _write_html(data, kind, src_ext, out_path):
    import pandas as pd
    if kind == "dataframe":
        body = data.to_html(index=False, border=1)
        title = os.path.basename(out_path)
    elif kind == "html_pair":
        # already HTML — just save raw
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(data[0])
        return
    elif src_ext in CODE_EXTS:
        lang = src_ext.lstrip(".")
        text = _to_plain_text(data, kind).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        body = f"<pre><code class='language-{lang}'>{text}</code></pre>"
        title = os.path.basename(out_path)
    else:
        text = _to_plain_text(data, kind).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        body = "<br>\n".join(f"<p>{line}</p>" for line in text.splitlines() if line.strip())
        title = os.path.basename(out_path)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>{title}</title>
<style>body{{font-family:sans-serif;padding:2em}}table{{border-collapse:collapse}}td,th{{padding:6px 12px;border:1px solid #ccc}}</style>
</head>
<body>
{body}
</body>
</html>"""
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)

def _write_pdf(data, kind, src_ext, out_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch

    styles = getSampleStyleSheet()
    story  = []

    if kind == "dataframe":
        # Table output
        df       = data.fillna("").astype(str)
        headers  = list(df.columns)
        rows     = [headers] + df.values.tolist()
        t        = Table(rows, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",  (0,0), (-1,0),  colors.HexColor("#4472C4")),
            ("TEXTCOLOR",   (0,0), (-1,0),  colors.white),
            ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EBF3FB")]),
            ("GRID",        (0,0), (-1,-1), 0.5, colors.grey),
            ("ALIGN",       (0,0), (-1,-1), "LEFT"),
            ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
            ("TOPPADDING",  (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ]))
        story.append(t)
    elif src_ext in CODE_EXTS:
        text = _to_plain_text(data, kind)
        story.append(Preformatted(text, styles["Code"]))
    else:
        text = _to_plain_text(data, kind)
        for line in text.splitlines():
            if line.strip():
                story.append(Paragraph(line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles["Normal"]))
            else:
                story.append(Spacer(1, 0.1*inch))

    doc = SimpleDocTemplate(out_path, pagesize=A4,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    doc.build(story)

def _write_docx(data, kind, src_ext, out_path):
    from docx import Document
    doc = Document()

    if kind == "dataframe":
        df    = data.fillna("").astype(str)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"
        hdr   = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    else:
        text = _to_plain_text(data, kind)
        for line in text.splitlines():
            doc.add_paragraph(line)

    doc.save(out_path)

def _write_csv(data, kind, src_ext, out_path):
    import pandas as pd
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        df.to_csv(out_path, index=False)
    else:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(_to_plain_text(data, kind))

def _write_tsv(data, kind, src_ext, out_path):
    import pandas as pd
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        df.to_csv(out_path, index=False, sep="\t")
    else:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(_to_plain_text(data, kind))

def _write_json(data, kind, src_ext, out_path):
    import pandas as pd
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        records = df.to_dict(orient="records")
    elif kind in ("json_obj", "yaml_obj"):
        records = data
    else:
        records = _to_plain_text(data, kind)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(records, f, indent=2, ensure_ascii=False)

def _write_xlsx(data, kind, src_ext, out_path):
    import pandas as pd
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        df.to_excel(out_path, index=False)
    else:
        text = _to_plain_text(data, kind)
        pd.DataFrame([[line] for line in text.splitlines()],
                     columns=["content"]).to_excel(out_path, index=False)

def _write_xml(data, kind, src_ext, out_path):
    import xml.etree.ElementTree as ET
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        root = ET.Element("data")
        for _, row in df.iterrows():
            item = ET.SubElement(root, "row")
            for col, val in row.items():
                child = ET.SubElement(item, str(col).replace(" ","_"))
                child.text = str(val)
        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ")
        tree.write(out_path, encoding="unicode", xml_declaration=True)
    elif kind == "xml_tree":
        root, tree = data
        ET.indent(tree, space="  ")
        tree.write(out_path, encoding="unicode", xml_declaration=True)
    else:
        root = ET.Element("content")
        root.text = _to_plain_text(data, kind)
        ET.ElementTree(root).write(out_path, encoding="unicode", xml_declaration=True)

def _write_yaml(data, kind, src_ext, out_path):
    import pandas as pd
    import yaml
    df = _to_dataframe(data, kind, src_ext)
    if df is not None:
        records = df.to_dict(orient="records")
    elif kind == "xml_tree":
        root, _ = data
        records = {root.tag: [{c.tag: c.text for c in child} for child in root]}
    elif kind in ("json_obj", "yaml_obj"):
        records = data
    else:
        records = _to_plain_text(data, kind)
    with open(out_path, "w", encoding="utf-8") as f:
        yaml.dump(records, f, allow_unicode=True, default_flow_style=False, sort_keys=False)


WRITERS = {
    ".txt":  _write_txt,
    ".md":   _write_md,
    ".html": _write_html,
    ".htm":  _write_html,
    ".pdf":  _write_pdf,
    ".docx": _write_docx,
    ".csv":  _write_csv,
    ".tsv":  _write_tsv,
    ".json": _write_json,
    ".xlsx": _write_xlsx,
    ".xml":  _write_xml,
    ".yaml": _write_yaml,
    ".yml":  _write_yaml,
}


# ══════════════════════════════════════════════════════════════════════════════
# PREVIEW
# ══════════════════════════════════════════════════════════════════════════════

def _preview(data, kind, target_ext, src_ext, lines=8):
    """Return a short preview string of what the output would look like."""
    try:
        if target_ext in (".csv", ".tsv"):
            df = _to_dataframe(data, kind, src_ext)
            if df is not None:
                sep = "\t" if target_ext == ".tsv" else ","
                return df.head(lines).to_csv(index=False, sep=sep).strip()
        if target_ext == ".json":
            df = _to_dataframe(data, kind, src_ext)
            if df is not None:
                return json.dumps(df.head(lines).to_dict(orient="records"), indent=2)
            if kind in ("json_obj", "yaml_obj"):
                return json.dumps(data, indent=2)[:500]
        if target_ext in (".txt", ".md"):
            text = _to_plain_text(data, kind)
            return "\n".join(text.splitlines()[:lines])
        if target_ext == ".xml":
            df = _to_dataframe(data, kind, src_ext)
            if df is not None:
                import xml.etree.ElementTree as ET
                root = ET.Element("data")
                for _, row in df.head(3).iterrows():
                    item = ET.SubElement(root, "row")
                    for col, val in row.items():
                        c = ET.SubElement(item, str(col).replace(" ","_"))
                        c.text = str(val)
                ET.indent(ET.ElementTree(root), space="  ")
                return ET.tostring(root, encoding="unicode")[:600]
        if target_ext in (".yaml", ".yml"):
            import yaml
            df = _to_dataframe(data, kind, src_ext)
            if df is not None:
                return yaml.dump(df.head(lines).to_dict(orient="records"),
                                 allow_unicode=True, default_flow_style=False)[:500]
        if target_ext == ".html":
            df = _to_dataframe(data, kind, src_ext)
            if df is not None:
                return df.head(lines).to_html(index=False, border=1)[:600]
    except Exception as e:
        return f"(Preview error: {e})"
    return "(Preview not available for this format)"


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

FORMAT_WORDS = {
    "csv": ".csv", "tsv": ".tsv", "json": ".json", "excel": ".xlsx",
    "xlsx": ".xlsx", "xls": ".xlsx", "xml": ".xml", "yaml": ".yaml",
    "yml": ".yaml", "text": ".txt", "txt": ".txt", "markdown": ".md",
    "md": ".md", "pdf": ".pdf", "word": ".docx", "docx": ".docx",
    "html": ".html", "htm": ".html",
}

def _parse_target_ext(instructions, src_ext):
    """
    Extract the TARGET format from the instruction string.
    Key fix: ignore any extension that matches the SOURCE file's own name/ext,
    since the instruction often contains the source filename (e.g. "convert merged.csv to txt").
    We search from RIGHT to LEFT so the target format mentioned last wins.
    """
    if not instructions:
        return None

    # Strip the source filename from the instruction to avoid false matches.
    # e.g. "convert merged.csv to json" → we remove "merged.csv" before searching
    src_pattern = r"\b[\w.-]+" + re.escape(src_ext) + r"\b"
    clean = re.sub(src_pattern, "", instructions, flags=re.IGNORECASE).strip()

    # 1. Explicit extension literal in the cleaned string
    for m in re.finditer(r"\.(csv|tsv|json|xlsx|xml|yaml|yml|txt|md|pdf|docx|html|htm)\b", clean, re.I):
        ext = "." + m.group(1).lower()
        ext = ".yaml" if ext == ".yml" else (".html" if ext == ".htm" else ext)
        if ext != src_ext:
            return ext  # return first match after source is stripped

    # 2. Format keyword match on cleaned string
    low = clean.lower()
    for word, ext in FORMAT_WORDS.items():
        if re.search(r"\b" + re.escape(word) + r"\b", low) and ext != src_ext:
            return ext

    return None

def _output_path(src_path, target_ext):
    base = os.path.splitext(src_path)[0]
    return base + target_ext


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def run(filename=None, instructions=None, **kwargs):

    # 1. Resolve source
    src = resolve_single(filename, instructions, prompt_label="file to convert")
    if not src:
        print("  No source file found.")
        return

    src_ext  = os.path.splitext(src)[1].lower()
    src_name = os.path.basename(src)

    if src_ext not in READERS:
        print(f"  ❌ Unsupported source format: '{src_ext}'")
        print(f"     Supported sources: {', '.join(sorted(READERS))}")
        return

    # 2. Determine target format
    target_ext = _parse_target_ext(instructions or "", src_ext)

    if not target_ext:
        allowed = sorted(_allowed_targets(src_ext))
        print(f"\n  Source : {src_name}  [{src_ext}]")
        cat = "data" if src_ext in DATA_EXTS else ("document" if src_ext in DOC_EXTS else "code/config")
        print(f"  Category: {cat}")
        print(f"  Can convert to: {', '.join(allowed)}")
        raw = input("  Target format: ").strip().lower().lstrip(".")
        if not raw:
            return
        target_ext = "." + raw
        target_ext = FORMAT_WORDS.get(raw, target_ext)

    # normalise aliases
    if target_ext == ".yml":
        target_ext = ".yaml"
    if target_ext == ".htm":
        target_ext = ".html"

    if target_ext == src_ext or (src_ext == ".yml" and target_ext == ".yaml"):
        print("  Source and target formats are the same — nothing to do.")
        return

    allowed = _allowed_targets(src_ext)
    if target_ext not in allowed:
        print(f"  ❌ Cannot convert {src_ext} → {target_ext}")
        print(f"     Allowed targets for {src_ext}: {', '.join(sorted(allowed))}")
        return

    if target_ext not in WRITERS:
        print(f"  ❌ No writer available for '{target_ext}'")
        return

    # 3. Output path
    out_path  = _output_path(src, target_ext)
    out_name  = os.path.basename(out_path)

    # 4. Read source
    print(f"\n  ┌─ Conversion Plan ─────────────────────────────")
    print(f"  │  Source  : {src_name}")
    print(f"  │  Output  : {out_name}")
    print(f"  │  Format  : {src_ext}  →  {target_ext}")
    print(f"  └───────────────────────────────────────────────")

    try:
        data, kind = READERS[src_ext](src)
    except Exception as e:
        print(f"\n  ❌ Read error: {e}")
        return

    # 5. Preview
    print(f"\n  ── Preview of output ({out_name}) ──────────────\n")
    prev = _preview(data, kind, target_ext, src_ext)
    for line in prev.splitlines()[:12]:
        print(f"  {line}")
    print(f"  ...")

    # 6. Handle existing output file
    if os.path.exists(out_path):
        print(f"\n  ⚠  '{out_name}' already exists.")
        choice = input("  1 → overwrite   2 → rename   3 → cancel: ").strip()
        if choice == "2":
            new_name = input("  New output filename: ").strip()
            out_path = os.path.join(os.path.dirname(src), new_name)
            out_name = os.path.basename(out_path)
        elif choice != "1":
            print("  Cancelled.")
            return

    ok, reason = run_write_guardrails(out_path)
    if not ok:
        print(f"  Blocked: {reason}")
        return

    if not confirm(f"\n  Convert '{src_name}' → '{out_name}'? (yes/no): "):
        print("  Cancelled.")
        return

    # 7. Write
    try:
        WRITERS[target_ext](data, kind, src_ext, out_path)
    except Exception as e:
        print(f"  ❌ Write error: {e}")
        import traceback; traceback.print_exc()
        return

    size = os.path.getsize(out_path)
    size_str = f"{size:,} B" if size < 1024 else f"{size//1024:,} KB"
    print(f"\n  ✅ Converted: {src_name}  →  {out_name}  ({size_str})")
